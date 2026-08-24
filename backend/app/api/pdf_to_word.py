from app.services.activity_service import record_activity
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse

import io
import os
import shutil
import tempfile

import pymupdf
from docx import Document
from docx.shared import Inches, Pt


router = APIRouter(
    prefix="/api",
    tags=["PDF to Word"],
)


MAX_FILE_SIZE = 20 * 1024 * 1024


def add_text_block(document, text, font_size=11):
    """
    Menambahkan blok teks ke Word.
    """
    text = text.strip()

    if not text:
        return

    paragraph = document.add_paragraph()

    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)

    paragraph.paragraph_format.space_after = Pt(4)


def add_image_block(document, image_bytes, page_width):
    """
    Menambahkan gambar PDF ke Word dengan ukuran proporsional.
    """
    try:
        image_stream = io.BytesIO(image_bytes)

        # Lebar maksimal gambar mengikuti lebar halaman Word.
        max_width = min(
            page_width,
            Inches(6.5),
        )

        document.add_picture(
            image_stream,
            width=max_width,
        )

    except Exception:
        pass


def convert_pdf_to_docx(input_path, output_path):
    """
    Konversi PDF ke DOCX dengan mempertahankan:
    - urutan blok
    - teks
    - gambar
    - pemisahan halaman
    """

    pdf = pymupdf.open(input_path)

    if pdf.page_count == 0:
        pdf.close()
        raise ValueError("PDF tidak memiliki halaman.")

    document = Document()

    # Margin Word.
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    page_width = (
        section.page_width
        - section.left_margin
        - section.right_margin
    )

    for page_number, page in enumerate(pdf):

        # -------------------------------------------------
        # Ambil semua blok berdasarkan posisi PDF
        # -------------------------------------------------

        blocks = page.get_text(
            "dict",
            flags=pymupdf.TEXTFLAGS_TEXT,
        ).get("blocks", [])

        # -------------------------------------------------
        # Urutkan berdasarkan posisi Y kemudian X
        # -------------------------------------------------

        text_blocks = []

        for block in blocks:

            if block.get("type") != 0:
                continue

            lines = block.get("lines", [])

            block_text = []

            font_sizes = []

            for line in lines:

                line_text = []

                for span in line.get("spans", []):

                    value = span.get(
                        "text",
                        "",
                    )

                    if value:
                        line_text.append(value)

                    size = span.get(
                        "size",
                        11,
                    )

                    if size:
                        font_sizes.append(size)

                if line_text:
                    block_text.append(
                        "".join(line_text)
                    )

            text = "\n".join(block_text).strip()

            if not text:
                continue

            bbox = block.get(
                "bbox",
                (0, 0, 0, 0),
            )

            average_size = (
                sum(font_sizes) / len(font_sizes)
                if font_sizes
                else 11
            )

            text_blocks.append(
                {
                    "type": "text",
                    "bbox": bbox,
                    "text": text,
                    "font_size": average_size,
                }
            )

        # -------------------------------------------------
        # Ambil gambar
        # -------------------------------------------------

        image_blocks = []

        for image_info in page.get_images(
            full=True
        ):

            xref = image_info[0]

            try:
                image_data = pdf.extract_image(
                    xref
                )

                if not image_data:
                    continue

                image_bytes = image_data.get(
                    "image"
                )

                if not image_bytes:
                    continue

                # Cari posisi gambar di halaman.
                rects = page.get_image_rects(
                    xref
                )

                if rects:

                    rect = rects[0]

                    image_blocks.append(
                        {
                            "type": "image",
                            "bbox": (
                                rect.x0,
                                rect.y0,
                                rect.x1,
                                rect.y1,
                            ),
                            "image": image_bytes,
                        }
                    )

            except Exception:
                continue

        # -------------------------------------------------
        # Gabungkan teks + gambar
        # berdasarkan posisi halaman
        # -------------------------------------------------

        elements = (
            text_blocks +
            image_blocks
        )

        elements.sort(
            key=lambda item: (
                item["bbox"][1],
                item["bbox"][0],
            )
        )

        # -------------------------------------------------
        # Masukkan ke Word
        # -------------------------------------------------

        for element in elements:

            if element["type"] == "text":

                text = element["text"]

                font_size = element.get(
                    "font_size",
                    11,
                )

                # Ukuran PDF sering menggunakan
                # nilai lebih besar daripada ukuran Word.
                font_size = max(
                    8,
                    min(
                        float(font_size),
                        28,
                    ),
                )

                add_text_block(
                    document,
                    text,
                    font_size,
                )

            elif element["type"] == "image":

                add_image_block(
                    document,
                    element["image"],
                    page_width,
                )

        # -------------------------------------------------
        # Page break
        # -------------------------------------------------

        if page_number < pdf.page_count:

            document.add_page_break()

    pdf.close()

    document.save(output_path)


@router.post("/pdf-to-word")
async def pdf_to_word(
    file: UploadFile = File(...),
):

    if not file.filename:

        raise HTTPException(
            status_code=400,
            detail="Nama file tidak ditemukan.",
        )

    if not file.filename.lower().endswith(
        ".pdf"
    ):

        raise HTTPException(
            status_code=400,
            detail="File harus berformat PDF.",
        )

    contents = await file.read()

    if not contents:

        raise HTTPException(
            status_code=400,
            detail="File PDF kosong.",
        )

    if len(contents) > MAX_FILE_SIZE:

        raise HTTPException(
            status_code=400,
            detail="Ukuran file maksimal 20 MB.",
        )

    temp_dir = tempfile.mkdtemp(
        prefix="pdf_to_word_"
    )

    input_path = os.path.join(
        temp_dir,
        "input.pdf",
    )

    output_path = os.path.join(
        temp_dir,
        "output.docx",
    )

    try:

        with open(
            input_path,
            "wb",
        ) as f:

            f.write(contents)

        convert_pdf_to_docx(
            input_path,
            output_path,
        )

        if not os.path.exists(
            output_path
        ):

            raise RuntimeError(
                "File Word gagal dibuat."
            )

        with open(
            output_path,
            "rb",
        ) as f:

            docx_data = f.read()

        if not docx_data:

            raise RuntimeError(
                "File Word kosong."
            )


        original_name = os.path.splitext(
            file.filename
        )[0]

        output_filename = (
            f"{original_name}.docx"
        )


        record_activity(
            user_id=1,
            activity_type="pdf-to-word",
            title="PDF to Word",
            filename=file.filename,
            status="success",
            file_size=len(docx_data),
            metadata={
                "original_size": len(contents),
                "output_size": len(docx_data),
                "output_filename": output_filename
            }
        )

        return StreamingResponse(
            io.BytesIO(docx_data),

            media_type=(
                "application/vnd.openxmlformats-"
                "officedocument.wordprocessingml.document"
            ),

            headers={
                "Content-Disposition": (
                    f'attachment; '
                    f'filename="{output_filename}"'
                ),

                "Access-Control-Expose-Headers": (
                    "X-Original-Size, "
                    "X-Output-Size"
                ),

                "X-Original-Size": str(
                    len(contents)
                ),

                "X-Output-Size": str(
                    len(docx_data)
                ),
            },
        )

    except HTTPException:
        raise

    except Exception as error:

        raise HTTPException(
            status_code=500,
            detail=(
                f"Gagal mengonversi PDF ke Word: "
                f"{str(error)}"
            ),
        )

    finally:

        shutil.rmtree(
            temp_dir,
            ignore_errors=True,
        )








